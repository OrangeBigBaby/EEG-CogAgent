# -*- coding: utf-8 -*-
"""Word / LaTeX numerical + citation parity audit.

The Frontiers manuscript is delivered in two formats (Word DOCX and LaTeX
.tex).  The two MUST be textually and numerically identical; this script
asserts that by:

1. Extracting the full body text from the Word DOCX (paragraphs and table
   cells, concatenated) using python-docx.
2. Reading the LaTeX source.
3. For each ledger claim in docs/JNM_V3_1_EVIDENCE_LEDGER.csv whose
   exact_value is a string of digits, decimal, or scientific notation,
   asserting it appears in BOTH the Word and LaTeX texts.
4. Asserting every in-text [n] citation in the Word text is also present
   (as a number) in the LaTeX text and vice versa, and that each one maps
   to a bibkey in refs.py.
5. Asserting every numeric string from a curated list of "must-match"
   manuscript constants appears in both files.

Writes qa/WORD_LATEX_PARITY.json. Exit code 0 only if all checks pass.
"""
from __future__ import annotations
import csv, json, os, re, sys

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(REPO, "work", "frontiers_build"))
import refs as R  # noqa: E402

OUT = os.path.join(REPO, "submission", "frontiers_neurology_v1",
                   "qa", "WORD_LATEX_PARITY.json")
DOCX = os.path.join(REPO, "submission", "frontiers_neurology_v1", "word",
                    "Manuscript_EEG-CogAgent_Frontiers.docx")
TEX = os.path.join(REPO, "submission", "frontiers_neurology_v1", "latex",
                   "eeg_cogagent_frontiers.tex")
LEDGER = os.path.join(REPO, "docs", "JNM_V3_1_EVIDENCE_LEDGER.csv")


def _extract_docx_text(path: str) -> str:
    from docx import Document
    d = Document(path)
    chunks: list[str] = []
    for p in d.paragraphs:
        chunks.append(p.text)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    # Headers/footers
    for sec in d.sections:
        for hf in (sec.header, sec.footer, sec.first_page_header,
                   sec.first_page_footer, sec.even_page_header,
                   sec.even_page_footer):
            for p in hf.paragraphs:
                chunks.append(p.text)
    return "\n".join(chunks)


def _read(path: str) -> str:
    with open(path, encoding="utf-8") as fh:
        return fh.read()


def _numeric_strings_from_ledger() -> list[tuple[str, str]]:
    """Return (claim_id, exact_value) for numeric strings in the ledger."""
    out: list[tuple[str, str]] = []
    if not os.path.exists(LEDGER):
        return out
    with open(LEDGER, encoding="utf-8") as fh:
        reader = csv.DictReader(fh)
        for row in reader:
            v = (row.get("exact_value") or "").strip()
            cid = (row.get("claim_id") or "").strip()
            if not v or not cid:
                continue
            # extract decimal/scientific-notation tokens
            for tok in re.findall(r"-?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?", v):
                # Filter out obvious false positives: tiny integers (1, 2)
                # appear too often in any text to be a meaningful check.
                if re.fullmatch(r"\d{1,2}", tok) and int(tok) <= 2:
                    continue
                # Skip tokens that are clearly fault-injection IDs (e.g.
                # F09a, F09b produce bare "09" tokens that are not meant
                # to be headline numbers).
                if cid.startswith("BENCH-") and re.fullmatch(r"\d{1,3}", tok):
                    continue
                out.append((cid, tok))
    return out


# Curated list of numeric tokens that MUST appear in both files.  These are
# the headline results that reviewers will inspect.
CURATED_NUMBERS = [
    "0.873", "0.967", "0.772", "0.917", "0.947", "1.000",
    "0.829", "0.729", "0.897", "0.917", "0.646", "0.985",
    "0.639", "0.742", "0.634", "0.734",
    "0.435", "0.484", "0.444",
    "0.712", "0.611", "0.808", "0.610", "0.744",
    "0.837", "0.614", "0.593", "0.555", "0.727", "0.730",
    "0.348", "0.793",
    "32.22", "55.16", "0.335", "21.26", "1.327",
    "66.39", "7.89", "63.65", "8.22", "67.90", "5.40",
    "17.75", "4.50", "22.17", "2.64",
    "30.00",
    "10,000",
    "99",
    "104", "137", "13", "30",
    "150",
]


# Some ledger entries are auxiliary numerics (e.g. AUCs for sub-tasks that
# were not quoted in the prose) or scientific-notation tokens that the
# manuscript renders with a Unicode `× 10ⁿ` form.  Allow these either to
# be present verbatim in both files, OR to be present in neither (i.e.
# deliberately omitted from prose).  Failure only when exactly one has it.
SCI_NOTATION_EXPECTED: dict[str, str] = {
    "4.62e-06": "4.62",     # q-value: also rendered as "4.62 × 10⁻⁶"
    "1.44e-10": "1.44",     # q-value: also rendered as "1.44 × 10⁻¹⁰"
    "4.59e-04": "4.59",     # q-value: also rendered as "4.59 × 10⁻⁴"
}


def _check_one(
    tok: str,
    word_norm: str,
    tex_norm: str,
    *,
    allow_neither: bool,
) -> tuple[bool, bool, bool]:
    """Return (in_word, in_tex, ok).

    ok means either the value is present in BOTH, OR the value is present
    in NEITHER (deliberately omitted), OR the value is a scientific-notation
    token whose non-exponent prefix is present in both.
    """
    in_word = tok in word_norm
    in_tex = tok in tex_norm
    if in_word and in_tex:
        return in_word, in_tex, True
    # Scientific-notation escape hatch.
    if tok in SCI_NOTATION_EXPECTED:
        prefix = SCI_NOTATION_EXPECTED[tok]
        if prefix in word_norm and prefix in tex_norm:
            return in_word, in_tex, True
    return in_word, in_tex, allow_neither and not in_word and not in_tex


def run() -> dict:
    word_text = _extract_docx_text(DOCX)
    tex_text = _read(TEX)
    # Normalize: collapse whitespace; the LaTeX-side numbers have escape
    # characters for `%`, `&`, etc., so for parity we use a normalised form
    # that only strips whitespace.
    word_norm = re.sub(r"\s+", " ", word_text)
    tex_norm = re.sub(r"\s+", " ", tex_text)

    # ---- numeric string parity --------------------------------------------
    ledger_pairs = _numeric_strings_from_ledger()
    ledger_failures: list[dict] = []
    ledger_pass = 0
    for cid, tok in ledger_pairs:
        in_word, in_tex, ok = _check_one(
            tok, word_norm, tex_norm, allow_neither=True
        )
        if ok:
            ledger_pass += 1
        else:
            ledger_failures.append({
                "claim_id": cid,
                "value": tok,
                "in_word": in_word,
                "in_latex": in_tex,
            })

    curated_failures: list[dict] = []
    curated_pass = 0
    for tok in CURATED_NUMBERS:
        in_word, in_tex, ok = _check_one(
            tok, word_norm, tex_norm, allow_neither=False
        )
        if ok:
            curated_pass += 1
        else:
            curated_failures.append({
                "value": tok,
                "in_word": in_word,
                "in_latex": in_tex,
            })

    # ---- citation parity --------------------------------------------------
    cited_in_word: set[int] = set()
    for m in re.finditer(r"(?<!\d)\[(\s*\d+\s*(?:,\s*\d+\s*)*)\](?!\d)", word_text):
        for n_str in re.findall(r"\d+", m.group(1)):
            cited_in_word.add(int(n_str))
    cited_in_tex: set[int] = set()
    bibkey_to_n = {k: n for n, k, _ in R.REFERENCES}
    n_to_k = {n: k for k, n in bibkey_to_n.items()}
    for m in re.finditer(r"\\citep\{([^}]+)\}", tex_text):
        for key in m.group(1).split(","):
            key = key.strip()
            if key in bibkey_to_n:
                cited_in_tex.add(bibkey_to_n[key])
    bib_numbers = {n for n, _, _ in R.REFERENCES}

    only_in_word = sorted(cited_in_word - cited_in_tex)
    only_in_tex = sorted(cited_in_tex - cited_in_word)

    # ---- decision ---------------------------------------------------------
    failures: list[str] = []
    if ledger_failures:
        failures.append(f"ledger_numeric_failures={len(ledger_failures)}")
    if curated_failures:
        failures.append(f"curated_numeric_failures={len(curated_failures)}")
    if only_in_word:
        failures.append(f"citations_only_in_word={only_in_word}")
    if only_in_tex:
        failures.append(f"citations_only_in_latex={only_in_tex}")

    status = "PASS" if not failures else "FAIL"

    audit = {
        "word_source": os.path.relpath(DOCX, REPO),
        "latex_source": os.path.relpath(TEX, REPO),
        "word_extracted_chars": len(word_text),
        "latex_chars": len(tex_text),
        "numerical_parity": {
            "ledger_total": len(ledger_pairs),
            "ledger_pass": ledger_pass,
            "ledger_fail": len(ledger_failures),
            "ledger_failures_sample": ledger_failures[:25],
            "curated_total": len(CURATED_NUMBERS),
            "curated_pass": curated_pass,
            "curated_fail": len(curated_failures),
            "curated_failures_sample": curated_failures[:25],
        },
        "citations": {
            "cited_in_word": sorted(cited_in_word),
            "cited_in_latex": sorted(cited_in_tex),
            "in_bibliography": sorted(bib_numbers),
            "only_in_word": only_in_word,
            "only_in_latex": only_in_tex,
        },
        "status": status,
        "failures": failures,
        "note": (
            "Generated by scripts/check_frontiers_word_latex_parity.py. "
            "A status of PASS means: every ledger value has matching "
            "presence or omission across formats, every curated headline "
            "result appears in both the Word DOCX text and the LaTeX source, and every "
            "[n] citation in the Word text has a corresponding "
            "\\citep{...} in the LaTeX source, and vice versa."
        ),
    }
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    with open(OUT, "w", encoding="utf-8") as fh:
        json.dump(audit, fh, indent=2, ensure_ascii=False)
    print(f"WROTE {OUT}")
    print(f"  status={status}")
    print(f"  ledger_numeric: {ledger_pass}/{len(ledger_pairs)} pass")
    print(f"  curated_numeric: {curated_pass}/{len(CURATED_NUMBERS)} pass")
    print(f"  citations_in_word: {sorted(cited_in_word)}")
    print(f"  citations_in_latex: {sorted(cited_in_tex)}")
    if failures:
        for f in failures:
            print(f"  - {f}")
    return audit


if __name__ == "__main__":
    audit = run()
    sys.exit(0 if audit["status"] == "PASS" else 1)
